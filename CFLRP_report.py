from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx2pdf import convert
import os
import datetime
import calendar
import matplotlib
matplotlib.use('Agg')
from matplotlib import pyplot as plt
from matplotlib.colors import ListedColormap
import numpy as np
import pandas as pd
import os
import arcpy
def updateTitleDoc(projectArea, projectsDict, pathToTitles, docxPath):
    ##update forest name, region, and date in word doc and convert to pdf to create first 2 pages of TCA report
    outputName = os.path.join(pathToTitles,'CFLRP_TitlePages_{}.docx'.format(projectArea))
    #if os.path.exists(outputName) == overwrite or os.path.exists(outputName) == False: 
    document = Document(docxPath)
    region = 'Region '+str(projectsDict[area][0])
    internal = document.paragraphs[2]
    txt = internal.text[:5]
    internal.text = ''
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('CommentsStyle3', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Arial'
    date = calendar.month_name[datetime.datetime.now().month] + ' '+str(datetime.datetime.now().year)
    internal.add_run(date, style = 'CommentsStyle3')
    section = document.paragraphs[5]       
    fontsize = 30 
    section.text = ''
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(fontsize)
    obj_font.name = 'Arial'
    section.add_run(projectArea, style = 'CommentsStyle')
    section2 = document.paragraphs[6]
    fontsize = 12
    section2.text = ''
    obj_charstyle2 = obj_styles.add_style('CommentsStyle2', WD_STYLE_TYPE.CHARACTER)
    obj_font2 = obj_charstyle2.font
    obj_font2.italic = True
    obj_font2.size = Pt(fontsize)
    obj_font2.name = 'Arial'
    section2.add_run(region, style = 'CommentsStyle2')
    if os.path.exists(outputName):
        os.remove(outputName)
    document.save(outputName)
    pdfOutputLoc = outputName[0:-4] + "pdf"
    print(f"Output location for pdf: {pdfOutputLoc}")
    print(f"Output location for word doc: {outputName}\n")
    print ('*** Do not open any docx files until all have been converted (this will overwrite current pdfs) *****\n')
    convert(outputName, pdfOutputLoc)
def feature_class_to_pandas_data_frame(feature_class, field_list):
    return pd.DataFrame(
        arcpy.da.FeatureClassToNumPyArray(
            in_table=feature_class,
            field_names=field_list,
            skip_nulls=False,
            null_value=-99999
        )
    )
def exportAndAppendToFinalPDF(srcLayout, tempPath, finalPath):
        srcLayout.exportToPDF(tempPath)#,resolution=500, output_as_image=True) #output as image to convert vector parts of pdf to raster (helps with width issue of LTA outline)
        finalPath.appendPages(tempPath)
def exportAndInsertToFinalPDF(srcLayout, tempPath, finalPath):
        srcLayout.exportToPDF(tempPath)#,resolution=500, output_as_image=True) #output as image to convert vector parts of pdf to raster (helps with width issue of LTA outline)
        finalPath.insertPages(tempPath,1)
def createNumberOfLandscapesTable(in1,in2):
    heading = "Number of Landscapes"
    size = (np.array(in1.shape[::-1]) + np.array([MPL_COL_WIDTH, 1])) * np.array([MPL_COL_WIDTH, MPL_ROW_HEIGHT])
    fig, ax = plt.subplots(figsize=size)
    ax.axis('off')
    mpl_table_header = ax.table(cellText=[['']],colLabels = [heading],loc='top')
    mpl_table = ax.table(cellText=in1.values, bbox=BBOX, colLabels=in1.columns,rowLabels= in1.index,loc='top')
    #ax.axhline(y=1, color=header_color, lw=1,alpha=1)  #alpha=1, 
    ax.axvline(x=0,color='w', alpha=1, lw=1)
    ax.axvline(x=0,color='w', alpha=1, lw=1)
    mpl_table.auto_set_font_size(False)
    mpl_table.set_fontsize(FONT_SIZE)
    mpl_table_header.auto_set_font_size(False)
    mpl_table_header.set_fontsize(FONT_SIZE)
    for k, cell in mpl_table._cells.items():
        cell.set_edgecolor(EDGE_COLOR)
        if k[0] == 0 or k[1] < HEADER_COLUMNS:
            cell.set_text_props(weight='bold', color='w')
            cell.set_facecolor(HEADER_COLOR)
            if k[0] == 0:
                cell.set_text_props(weight='bold', color='w')#, va ='top')  
        else:
            cell.set_facecolor(ROW_COLORS[k[0]%len(ROW_COLORS) ])
        if k ==(1,-1):
            cell.set_facecolor("#B61D36")
        if k ==(2,-1):
            cell.set_facecolor("#FF9966")
        if k ==(3,-1):
            cell.set_facecolor("#DFD696")
        if k ==(4,-1):
            cell.set_facecolor("#00A3E6")
        if k ==(5,-1):
            cell.set_facecolor("#003399")
    for k, cell in mpl_table_header._cells.items():
        cell.set_edgecolor(EDGE_COLOR)
        if k[0] == 0 or k[1] < HEADER_COLUMNS:           
            cell.set_facecolor(HEADER_COLOR)
            #h = cell.get_height()
            #cell.set_height(h/1.5)
            cell.set_text_props(weight='bold', color='w')#, verticalalignment ='bottom')
    #adjust height of first 'row' of initial table
    cellDict = mpl_table_header.get_celld()
    cellDict[(1,0)].set_height(0)
    if len(heading) > 25:
        cellDict[(0,0)].set_height(1/3.8)#6.) 
    else:
        cellDict[(0,0)].set_height(1/6.) 
    outName = in2 + "_table_byClass"
    outFile = os.path.join(OUTPUT_PATH, outName)
    plt.savefig(outFile, dpi=300,bbox_inches=BBOX_INCHES)
    reg_countTable = outFile + ".png"
    plt.close()
def createAcresTable(in1,in2):
    heading = "Area (in thoudsands of acres)"
    size = (np.array(in1.shape[::-1]) + np.array([MPL_COL_WIDTH, 1])) * np.array([MPL_COL_WIDTH, MPL_ROW_HEIGHT])
    fig, ax = plt.subplots(figsize=size)
    ax.axis('off')
    mpl_table_header = ax.table(cellText=[['']],colLabels = [heading],loc='top')
    mpl_table = ax.table(cellText=in1.values, bbox=BBOX, colLabels=in1.columns,rowLabels= in1.index,loc='top')
    #ax.axhline(y=1, color=header_color, lw=1,alpha=1)  #alpha=1, 
    ax.axvline(x=0,color='w', alpha=1, lw=1)
    ax.axvline(x=0,color='w', alpha=1, lw=1)
    mpl_table.auto_set_font_size(False)
    mpl_table.set_fontsize(FONT_SIZE)
    mpl_table_header.auto_set_font_size(False)
    mpl_table_header.set_fontsize(FONT_SIZE)
    for k, cell in mpl_table._cells.items():
        cell.set_edgecolor(EDGE_COLOR)
        if k[0] == 0 or k[1] < HEADER_COLUMNS:
            cell.set_text_props(weight='bold', color='w')
            cell.set_facecolor(HEADER_COLOR)
            if k[0] == 0:
                cell.set_text_props(weight='bold', color='w')#, va ='top')  
        else:
            cell.set_facecolor(ROW_COLORS[k[0]%len(ROW_COLORS) ])
        if k ==(1,-1):
            cell.set_facecolor("#B61D36")
        if k ==(2,-1):
            cell.set_facecolor("#FF9966")
        if k ==(3,-1):
            cell.set_facecolor("#DFD696")
        if k ==(4,-1):
            cell.set_facecolor("#00A3E6")
        if k ==(5,-1):
            cell.set_facecolor("#003399")
    for k, cell in mpl_table_header._cells.items():
        cell.set_edgecolor(EDGE_COLOR)
        if k[0] == 0 or k[1] < HEADER_COLUMNS:           
            cell.set_facecolor(HEADER_COLOR)
            #h = cell.get_height()
            #cell.set_height(h/1.5)
            cell.set_text_props(weight='bold', color='w')#, verticalalignment ='bottom')
    #adjust height of first 'row' of initial table
    cellDict = mpl_table_header.get_celld()
    cellDict[(1,0)].set_height(0)
    if len(heading) > 25:
        cellDict[(0,0)].set_height(1/3.8)#6.) 
    else:
        cellDict[(0,0)].set_height(1/6.) 
    outName = in2 + "_table_byClass_acres"
    outFile = os.path.join(OUTPUT_PATH, outName)
    plt.savefig(outFile, dpi=300,bbox_inches=BBOX_INCHES)
    reg_acresTable = outFile + ".png"
    plt.close()
def createBarChart(in1,in2):
    fig, ax = plt.subplots(1, 1)
    fig.set_size_inches(10,10)
    colors = ["#B61D36", "#FF9966", "#DFD696", "#00A3E6", "#003399"]
    custom_cmap = ListedColormap(colors)
    in1.transpose().plot(kind='barh', stacked=True, ax=ax, colormap=custom_cmap)
    ax.get_legend().remove()
    ax.spines['right'].set_visible(False)
    ax.spines['top'].set_visible(False)
    plt.xlabel('Number of LTAs',fontsize=32)
    plt.xticks(rotation=45, ha = 'right',size=24)
    plt.yticks(size=24)
    plt.ylabel('Reporting Year',fontsize=32)
    plt.tight_layout()
    outName = in2 + "_barchart_byClass"
    outFile = os.path.join(OUTPUT_PATH, outName)
    plt.savefig(outFile, dpi=300)
    reg_barChart = outFile + ".png"
    plt.close()
GDB_PATH = r"C:\Users\markhammond\Desktop\TCA_CFLRP_Databases.gdb"
arcpy.env.workspace = GDB_PATH
fcs = arcpy.ListFeatureClasses("TCA_Baseline*")
fc_path_list = [ ]
for fc in fcs:
    fc_path = os.path.join(GDB_PATH , fc)
    fc_path_list.append(fc_path)
YEAR_LIST = ['2020', '2021', '2022', '2023']
OUTPUT_PATH = r"C:\Users\markhammond\Desktop\cflrp\output" # outpath for figures
ANCILLARY_FILES_PATH = r"C:\Users\markhammond\Desktop\cflrp\Templates"
APRX = arcpy.mp.ArcGISProject("current")
INDICATOR_SYMBOLOGY_LAYER_PATH = os.path.join(ANCILLARY_FILES_PATH,'indicator_symbology.lyrx')
CFLRP_BOUNDARY_SYMBOLOGY_LAYER_PATH = os.path.join(ANCILLARY_FILES_PATH,'cflrp_boundary.lyrx')
cities = arcpy.mp.LayerFile(os.path.join(ANCILLARY_FILES_PATH,'USA Major Cities.lyrx'))
CFLRP_Boundaries_FC = "CFLRP_Boundaries"
Sections_FC = "Sections"
FINAL_OUTPUT_PATH = os.path.join(OUTPUT_PATH, "FinalOutput")
TEMP_PDF_PATH = os.path.join(OUTPUT_PATH, "temp.pdf")
OUTPUT_TITLE_PAGES_PATH = os.path.join(OUTPUT_PATH, "TitlePages")
APPENDIX_PDF_PATH = os.path.join(ANCILLARY_FILES_PATH, "CONUS_Appendix_Table_TEMPLATE.pdf")
GLOSSARY_PDF_PATH = os.path.join(ANCILLARY_FILES_PATH, "TCA_Glossary.pdf")
Regional_Terrestrial_Condition_Map = "region"
#projectsDict = {"Southern Blues Restoration Coalition":["6","TITLE_PAGES_TEMPLATE_SBRC.docx"],"Rio Chama Landscape":["3","TITLE_PAGES_TEMPLATE.docx"],"Southwest Colorado Collaborative Forest Landscape":["2","TITLE_PAGES_TEMPLATE.docx"]}
projectsDict = {"Southern Blues Restoration Coalition":["6","TITLE_PAGES_TEMPLATE_SBRC.docx"]}
cflrp_project_areas = list(projectsDict.keys())   
layout_text = {'M_tree_mortality': ['Tree Mortality', 'Significant tree mortality caused by insects and pathogen outbreaks','Forested Systems'], 
               'M_fire_deficit':['Ecological Process Departure','Fire deficit due to fire exclusion','All Systems'],
               "M_terrestrial_condition":["Overall Terrestrial Condition",'Overall score of ecological conditions synthesized from all indicators','All Systems'],
               "M_uncharacteristic_disturbance":["Uncharacteristic Disturbance Events",'Uncharacteristically severe wildfires','All Systems'],
               "M_wildfire_hazard_potential":["Wildfire Hazard Potential",'Significant fuel build-up','All Systems'],
               "M_insect_pathogen_hazard":["Insect and Disease Pathogen Hazard",'Susceptibility to insects and pathogens due to unhealthy forest conditions','Forested Systems'],
               "M_vegetation_departure":["Vegetation Departure",'Composition of vegetation or successional stages differ from the natural range of variation','Forested Systems']}
# include only certain fields when reading in the data for testing and efficiency
include = ["tca_id","tca_acres", "cflrp_project_name", "M_terrestrial_condition"]
inds = ["M_terrestrial_condition"]
#include = ["tca_id","tca_acres", "cflrp_project_name", "M_terrestrial_condition","M_uncharacteristic_disturbance","M_wildfire_hazard_potential","M_fire_deficit","M_tree_mortality","M_insect_pathogen_hazard","M_vegetation_departure"]
#inds = ["M_terrestrial_condition","M_uncharacteristic_disturbance","M_wildfire_hazard_potential","M_fire_deficit","M_tree_mortality","M_insect_pathogen_hazard","M_vegetation_departure"]
#table variables
HEADER_COLOR = '#00008B'
ROW_COLORS = ['#f1f1f2', 'w']
EDGE_COLOR='w'
BBOX = [0, 0, 1, 1]
BBOX_INCHES = 'tight'
FONT_SIZE = 18
HEADER_COLUMNS = 0
MPL_COL_WIDTH = 1.3
MPL_ROW_HEIGHT = 0.5
MPL_FINAL_COL_WIDTH = 1.9
MPL_FINAL_ROW_HEIGHT = 0.7
COLOR_ARR = np.array([
[0,0,139],
[115,178,255],
[245,202,122],
[255,167,127],
[255,0,0],
[211,211,211]])
for area in cflrp_project_areas:
    regional_acres_results = []
    regional_count_results = []
    PDF_path = os.path.join(FINAL_OUTPUT_PATH,"CFLRP_{}_Summary.pdf".format(area)) 
    finalPDF = arcpy.mp.PDFDocumentCreate(PDF_path)
    for fc in fc_path_list:  
        query_project = 'PROJECTNAME' + " = " + "'"+area+"'"   
        projectBoundary = arcpy.management.MakeFeatureLayer(CFLRP_Boundaries_FC, "project_boundary", query_project).getOutput(0)
        regionBoundary = arcpy.management.SelectLayerByLocation(Sections_FC, "INTERSECT", projectBoundary, "", "NEW_SELECTION") 
        arcpy.management.CopyFeatures(regionBoundary,"regionBoundary")
        regionBoundaryLayer = arcpy.management.MakeFeatureLayer("regionBoundary", "regionBoundaryLayer").getOutput(0)
        regionLTA = arcpy.management.SelectLayerByLocation(fc, "INTERSECT", "regionBoundaryLayer", "", "NEW_SELECTION")
        arcpy.management.CopyFeatures(regionLTA,"regionLTA")
        regionLayer = arcpy.management.MakeFeatureLayer("regionLTA", "regionLTA_layer").getOutput(0)
        projectLTA = arcpy.management.SelectLayerByLocation(fc, "INTERSECT", projectBoundary, "", "NEW_SELECTION")
        arcpy.management.CopyFeatures(projectLTA,"projectLTA_2023")
        projectLayer = arcpy.management.MakeFeatureLayer("projectLTA_2023", "projectLTA_layer_2023").getOutput(0)
        symbologyFields = "VALUE_FIELD {} {}".format("M_terrestrial_condition","M_terrestrial_condition")
        arcpy.management.ApplySymbologyFromLayer(projectLayer, INDICATOR_SYMBOLOGY_LAYER_PATH, symbologyFields)
        arcpy.management.ApplySymbologyFromLayer(regionLayer, INDICATOR_SYMBOLOGY_LAYER_PATH, symbologyFields)
        left_map = APRX.listMaps(Regional_Terrestrial_Condition_Map)[0]
        right_map = APRX.listMaps("2023")[0]                         
        left_map.addLayer(regionLayer)                         
        arcpy.management.ApplySymbologyFromLayer(projectBoundary,CFLRP_BOUNDARY_SYMBOLOGY_LAYER_PATH)
        left_map.addLayer(projectBoundary)
        right_map.addLayer(projectLayer)
        right_map.addLayer(projectBoundary)
        df = feature_class_to_pandas_data_frame(regionLayer, include)
        bins = [-1,-0.6,-0.2,0.2,0.6,1]
        labels=["Very Poor", "Poor", "Moderate", "Good", "Very Good"]
        df['Class'] = pd.cut(df["M_terrestrial_condition"], bins=bins, labels=labels, include_lowest=True)
        #get count by category
        count_by_class = df.groupby(['Class']).size()
        count_by_class.index.names = ["M_terrestrial_condition" + '_class']
        print(count_by_class)
        #get acres by category
        acres_by_class = df.groupby(['Class'])[include[1]].sum()/1000
        acres = acres_by_class.astype(int)
        acres.index.names = ["M_terrestrial_condition" + '_class']
        print(acres)
        regional_acres_results.append(acres)
        regional_count_results.append(count_by_class)
    reg_count_df = pd.concat([regional_count_results[0], regional_count_results[1], regional_count_results[2],regional_count_results[3]], axis = 1).rename(columns = {0:"2020",1:"2021",2:"2022",3:"2023"})
    print(reg_count_df)
    reg_acres_df = pd.concat([regional_acres_results[0], regional_acres_results[1], regional_acres_results[2],regional_acres_results[3]], axis = 1)
    reg_acres_df.columns = YEAR_LIST
    print(reg_acres_df)
    createNumberOfLandscapesTable(reg_count_df,"regional")
    createAcresTable(reg_acres_df,"regional")    
    createBarChart(reg_count_df,"regional")
    #update regional layout
    lay = APRX.listLayouts("Regional Terrestrial Condition")[0]
    pictures = lay.listElements("picture_element")
    for pic in pictures:
        if pic.name =='Picture 2':
            pic.sourceImage = reg_countTable
        elif pic.name =='Picture 1':
            pic.sourceImage = reg_acresTable
        elif pic.name =='Picture 3':
            pic.sourceImage = reg_barChart
    left_map.addLayer(cities)
    mf = lay.listElements('MAPFRAME_ELEMENT', 'Map Frame')[0]
    mf1 = lay.listElements('MAPFRAME_ELEMENT', 'Map Frame 1')[0]
    mf.camera.setExtent(mf.getLayerExtent(regionLayer))
    mf1.camera.setExtent(mf1.getLayerExtent(projectLayer))
    exportAndAppendToFinalPDF(lay, TEMP_PDF_PATH, finalPDF)                    
    for ind in inds:
        # initalize an empty lists
        acres_results = []
        count_results = []
        for fc in fc_path_list:  
            year = fc[-4:]
            current_map = APRX.listMaps(year)[0]
            projectBoundary = arcpy.management.MakeFeatureLayer("CFLRP_Boundaries", "project_bound", query_project).getOutput(0)
            arcpy.management.ApplySymbologyFromLayer(projectBoundary,CFLRP_BOUNDARY_SYMBOLOGY_LAYER_PATH)
            current_map.addLayer(projectBoundary)
            projectLTA = arcpy.management.SelectLayerByLocation(fc, "INTERSECT", projectBoundary, "", "NEW_SELECTION")
            arcpy.management.CopyFeatures(projectLTA,"projectLTA_"+ year)
            projectLayer = arcpy.management.MakeFeatureLayer("projectLTA_"+year, "projectLTA_layer_"+ year).getOutput(0)
            symbologyFields = "VALUE_FIELD {} {}".format(ind,ind)
            arcpy.management.ApplySymbologyFromLayer(projectLayer, INDICATOR_SYMBOLOGY_LAYER_PATH, symbologyFields)
            current_map.addLayer(projectLayer)
            current_map.addLayer(projectBoundary)
            df = feature_class_to_pandas_data_frame(projectLayer, include)
            bins = [-1,-0.6,-0.2,0.2,0.6,1]
            labels=["Very Poor", "Poor", "Moderate", "Good", "Very Good"]
            df['Class'] = pd.cut(df[ind], bins=bins, labels=labels, include_lowest=True)
            #get count by category
            count_by_class = df.groupby(['Class']).size()
            count_by_class.index.names = [ind + '_class']
            print(count_by_class)
            #get acres by category
            acres_by_class = df.groupby(['Class'])[include[1]].sum()/1000
            acres = acres_by_class.astype(int)
            acres.index.names = [ind + '_class']
            print(acres)
            acres_results.append(acres)
            count_results.append(count_by_class)
        lyt = APRX.listLayouts('Terrestrial Condition 2020-2023')[0]
        mf = lyt.listElements('MAPFRAME_ELEMENT', 'Map Frame')[0]
        mf1 = lyt.listElements('MAPFRAME_ELEMENT', 'Map Frame 1')[0]
        mf2 = lyt.listElements('MAPFRAME_ELEMENT', 'Map Frame 2')[0]
        mf3 = lyt.listElements('MAPFRAME_ELEMENT', 'Map Frame 3')[0]
        extent = arcpy.Describe(projectLayer).extent
        mf.camera.setExtent(extent)
        mf1.camera.setExtent(extent)
        mf2.camera.setExtent(extent)
        mf3.camera.setExtent(extent)
        count_df = pd.concat([count_results[0], count_results[1], count_results[2],count_results[3]], axis = 1).rename(columns = {0:"2020",1:"2021",2:"2022",3:"2023"})
        print(count_df)
        acres_df = pd.concat([acres_results[0], acres_results[1], acres_results[2],acres_results[3]], axis = 1)
        acres_df.columns = YEAR_LIST
        print(acres_df)
        createNumberOfLandscapesTable(count_df,ind)
        createAcresTable(acres_df,ind)    
        createBarChart(count_df,ind)
        #update indicator layout
        lay = APRX.listLayouts("Terrestrial Condition 2020-2023")[0]
        text = lay.listElements("text_element")
        pictures = lay.listElements("picture_element")
        for elm in text:
            if elm.name =='Text 7':
                elm.text = cflrp_project_areas[0]
            elif elm.name =='Text':
                elm.text = layout_text[ind][0]
            elif elm.name =='Text 2':
                elm.text = layout_text[ind][1]
            elif elm.name =='Text 10':
                elm.text = layout_text[ind][2]
        for pic in pictures:
            if pic.name =='Picture 2':
                pic.sourceImage = acresTable
            elif pic.name =='Picture 1':
                pic.sourceImage = barChart
            elif pic.name =='Picture 3':
                pic.sourceImage = countTable
        exportAndAppendToFinalPDF(lay, TEMP_PDF_PATH, finalPDF)
    for mp in APRX.listMaps():
        for rmlyr in mp.listLayers():
            if rmlyr.name != "World Ocean Base":
                mp.removeLayer(rmlyr)
 

    titlePage = os.path.join(OUTPUT_TITLE_PAGES_PATH, 'CFLRP_TitlePages_'+area+".pdf")
    TITLE_PAGES_TEMPLATE = os.path.join(ANCILLARY_FILES_PATH, projectsDict[area][1])
    updateTitleDoc(area, projectsDict, OUTPUT_TITLE_PAGES_PATH, TITLE_PAGES_TEMPLATE)
    finalPDF.insertPages(titlePage,1)
    finalPDF.appendPages(APPENDIX_PDF_PATH)
    finalPDF.appendPages(GLOSSARY_PDF_PATH)
    finalPDF.saveAndClose()
