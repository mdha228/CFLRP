"""
Run this script first for TCA PDF Report Production.
creates for each forest:
1. title page (based on existing document - just updates name of forest, current date, and region) --> 
    i.e. user needs to ensure that the main map in the title page is updated based on current TCA data 
            --> use title_page_map.mxd and set the data source of the top rendered layer to the latest dataset
            --> export map as png; clip down to bbox and insert into TCA_PDF_TITLE_PAGES_TEMPLATE.docx
2. per-forest tables for each indicator
3. per-forest bar plots for each indicator
4. final summary table
"""

def GENERATE_CONUS_REPORTS(INPUT_GDB_NAME, INPUT_GDB_FEATURE_CLASS_NAME, FOREST_DISSOLVE_LAYER_FILE_NAME, OUTPUT_DIRECTORY_NAME, REPORTING_YEAR, CWD):
    
    ####### IMPORT MODULES #######
    from docx import Document #to install, navigate to current python environement in cmd and run "pip install python-docx"
    from docx.shared import Pt
    from docx.enum.style import WD_STYLE_TYPE
    from docx2pdf import convert
    import os
    import datetime
    import calendar
    import matplotlib
    matplotlib.use('Agg')
    from matplotlib import pyplot as plt
    import numpy as np
    import arcpy
    import pandas as pd
    import gc
    import datetime

    ### CONUS EMDS MODEL VARIABLES
    from VARIABLES.CONUS_EMDS_MODEL import CONUS_EMDS_INDICATORS_FOR_REPORTS
    from VARIABLES.CONUS_EMDS_MODEL import CONUS_EMDS_CLIMATE_INDICATORS
    from VARIABLES.CONUS_EMDS_MODEL import CONUS_EMDS_INDICTORS_SUBTEXT
    from VARIABLES.CONUS_EMDS_MODEL import CONUS_CLIMATE_INDICATOR_CATEGORIES
    from VARIABLES.CONUS_EMDS_MODEL import CONUS_CLIMATE_APPENDIX_PAGES
    from VARIABLES.CONUS_EMDS_MODEL import CONUS_FIGURE_PATH_DICT
    from VARIABLES.CONUS_EMDS_MODEL import CONUS_FORESTED_INDICATOR_HEADERS
    from VARIABLES.CONUS_EMDS_MODEL import CONUS_EMDS_INDICATOR_TREE_ROOT

    ### CONUS FOREST VARIABLES
    from VARIABLES.FORESTS import CONUS_FORESTS_AND_REGIONS
    from VARIABLES.FORESTS import CONUS_FORESTS_WITH_NO_FOREST_LANDTYPE
    from VARIABLES.FORESTS import CONUS_FORESTS_WITH_NO_NON_FORESTED_LANDTYPES

    ### CONUS LTA-RELATED VARIABLES
    from VARIABLES.LTA_VARIABLES import EXTRA_CONUS_FIELDS

    ######## RELEVANT FILE PATHS #########

    ANCILLARY_FILES_PATH = os.path.join(CWD, "AncillaryFiles")

    INPUT_GDB_PATH = os.path.join(CWD, OUTPUT_DIRECTORY_NAME, INPUT_GDB_NAME)

    INPUT_GDB_FEATURE_CLASS_PATH = os.path.join(INPUT_GDB_PATH, INPUT_GDB_FEATURE_CLASS_NAME)

    DISSOLVED_FOREST_LAYER_PATH = os.path.join(CWD, OUTPUT_DIRECTORY_NAME, INPUT_GDB_NAME, FOREST_DISSOLVE_LAYER_FILE_NAME) # this is now a feature class instead of a lyrx file

    ### input template title page word doc
    TITLE_PAGES_TEMPLATE = os.path.join(ANCILLARY_FILES_PATH, 'TITLE_PAGE_TEMPLATE.docx')

    ### path of directory that is generated automatically for all outputs and reports.
    FINAL_OUTPUT_PATH = os.path.join(CWD, OUTPUT_DIRECTORY_NAME, "CFLRP_PROJECT_REPORTS")

    ### path for all generated title pages per forest.
    OUTPUT_TITLE_PAGES_PATH = os.path.join(FINAL_OUTPUT_PATH, "TitlePages")

    APPENDIX_PDF_PATH = os.path.join(ANCILLARY_FILES_PATH, "CONUS_Appendix_Table.pdf")

    GLOSSARY_PDF_PATH = os.path.join(ANCILLARY_FILES_PATH, "TCA_Glossary_CONUS.pdf")

    ### access symbology layer
    SYMBOLOGY_LAYER_FINAL_PATH = os.path.join(ANCILLARY_FILES_PATH,'Final_Symbology_20230425.lyrx')   # had to recreate this because field names are now different - no longer nw0__terrestrial_condition 

    ### symbology layer for forests
    SYMBOLOGY_LAYER_FORESTS_PATH = os.path.join(ANCILLARY_FILES_PATH,'ForestsLayer_fromTCAGDB_symbology5.lyrx')

    ### access mxd file called Lab Data
    MAPBOOK_PATH = os.path.join(ANCILLARY_FILES_PATH, "TCA_Mapbook_pro")

    #APRX_PATH = os.path.join(MAPBOOK_PATH, "TCA_Mapbook_pro_v13_OceanBasemap_Credits5_wm.aprx") 
    APRX_PATH = os.path.join(MAPBOOK_PATH, "TCA_CONUS_Indicator_Report_Layout.aprx") 

    #APRX_FINAL_PAGE_PATH = os.path.join(MAPBOOK_PATH, "TCA_Mapbook_pro_v9_TableFinalPage6.aprx") 
    APRX_FINAL_PAGE_PATH = os.path.join(MAPBOOK_PATH, "Final_Table_Page_CONUS.aprx") 

    #DROUGHT_APRX_PATH = os.path.join(MAPBOOK_PATH, "TCA_Mapbook_pro_Drought2.aprx") 
    DROUGHT_APRX_PATH = os.path.join(MAPBOOK_PATH, "Drought_Page.aprx") 

    #CLIMATE_FULL_APRX_PATH = os.path.join(MAPBOOK_PATH, "TCA_Mapbook_pro_SeasonalClimate_v2.aprx") 
    CLIMATE_FULL_APRX_PATH = os.path.join(MAPBOOK_PATH, "Seasonal_Climate_Full_Page_CONUS.aprx") 

    #CLIMATE_NO_OVERALL_APRX_PATH = os.path.join(MAPBOOK_PATH, "TCA_Mapbook_pro_SeasonalClimate_noOverall.aprx") 
    CLIMATE_NO_OVERALL_APRX_PATH = os.path.join(MAPBOOK_PATH, "Seasonal_Climate_No_Overall.aprx") 

    USDA_LOGO_PATH = os.path.join(ANCILLARY_FILES_PATH,"usda_logo_2.png")

    TEMP_PDF_PATH = os.path.join(FINAL_OUTPUT_PATH, 'temp.pdf')

    SYMBOLOGY_LAYER_PATH = os.path.join(CWD, "ForestsLayer_fromTCAGDB_subset.lyrx")

    TEST_LAYER_PATH = os.path.join(INPUT_GDB_PATH, "test_lyr.lyr")

    ######## VARIABLE DEFINITIONS ##########
    
    FS_REGIONS = [1,2,3,4,5,6,8,9] #Excludes Alaska (Region 10)

    INDICATOR_KEYS_LIST = list(CONUS_EMDS_INDICATORS_FOR_REPORTS.keys())

    FOREST_KEYS_LIST = list(CONUS_FORESTS_AND_REGIONS.keys())

    FIELD_NAMES_IN_GDB = [i.name for i in arcpy.ListFields(INPUT_GDB_FEATURE_CLASS_PATH)]

    SPATIAL_SUB_COLUMNS = ['Forest','Region','Nation']

    SEASONAL_SUB_COLUMNS = ['Spring','Summer','Fall', 'Winter']

    TCA_CLASSIFICATIONS = ['Very Good','Good','Moderate','Poor','Very Poor']

    LANDSCAPE_THOUSAND_ACRES_TABLE_TITLE = 'Landscape\n(in thousands of acres)'

    NUMBER_OF_LANDSCAPES_TABLE_TITLE = 'Number of Landscapes'

    HEADER_COLOR = '#00008B'

    ROW_COLORS = ['#f1f1f2', 'w']

    EDGE_COLOR='w'

    BBOX = [0, 0, 1, 1]

    BBOX_INCHES = 'tight'

    FONT_SIZE = 18

    HEADER_COLUMNS = 0

    MPL_COL_WIDTH = 1.3

    MPL_ROW_HEIGHT = 0.5

    MPL_FINAL_COL_WIDTH = 1.9

    MPL_FINAL_ROW_HEIGHT = 0.7

    COLOR_ARR = np.array([
                [0,0,139],
                [115,178,255],
                [245,202,122],
                [255,167,127],
                [255,0,0],
                [211,211,211]])

    ### text to tag pdf with. 
    TCA_VERSION = 'TCA.v2.' + str(REPORTING_YEAR)

    #define workspace as folder path
    arcpy.env.workspace = FINAL_OUTPUT_PATH

    #allow overwriting output files
    arcpy.env.overwriteOutput = True

    ### ARCGISPROJECT OBJECTS
    APRX = arcpy.mp.ArcGISProject(APRX_PATH)
    APRX_FINAL_PAGE = arcpy.mp.ArcGISProject(APRX_FINAL_PAGE_PATH)
    DROUGHT_APRX = arcpy.mp.ArcGISProject(DROUGHT_APRX_PATH)
    CLIMATE_FULL_APRX = arcpy.mp.ArcGISProject(CLIMATE_FULL_APRX_PATH)
    CLIMATE_NO_OVERALL_APRX = arcpy.mp.ArcGISProject(CLIMATE_NO_OVERALL_APRX_PATH)

    #APRX-related maps
    TOP_MAP = APRX.listMaps('TopPane')[0]

    ### APRX first layout, and lists of picture, text, and graphic elements
    APRX_LAYOUT_0 = APRX.listLayouts()[0]
    APRX_LAYOUT_0_PICTURE_ELEMENTS= APRX_LAYOUT_0.listElements("picture_element")
    APRX_LAYOUT_0_TEXT_ELEMENTS = APRX_LAYOUT_0.listElements("text_element")
    APRX_LAYOUT_0_GRAPHIC_ELEMENTS = APRX_LAYOUT_0.listElements('GRAPHIC_ELEMENT')

    ### APRX first layout, and lists of picture, text, and graphic elements
    APRX_LAYOUT_0_MAPFRAME_0 = APRX_LAYOUT_0.listElements("mapframe_element")[0] 
    #APRX_LAYOUT_0_MAPFRAME_1 = APRX_LAYOUT_0.listElements("mapframe_element")[1]

    ### APRX final page layout, and list of picture elements
    APRX_FINAL_PAGE_LAYOUT_0 = APRX_FINAL_PAGE.listLayouts()[0]
    APRX_FINAL_PAGE_LAYOUT_0_PICTURE_ELEMENTS = APRX_FINAL_PAGE_LAYOUT_0.listElements("picture_element")

    ### Drought APRX layout and lists of mapframe, picture, and text elements
    DROUGHT_APRX_LAYOUT_0 = DROUGHT_APRX.listLayouts()[0]
    DROUGHT_APRX_MAPFRAME_0 = DROUGHT_APRX_LAYOUT_0.listElements("mapframe_element")[0]
    DROUGHT_APRX_LAYOUT_0_PICTURE_ELEMENTS= DROUGHT_APRX_LAYOUT_0.listElements("picture_element")
    DROUGHT_APRX_LAYOUT_0_TEXT_ELEMENTS = DROUGHT_APRX_LAYOUT_0.listElements("text_element")

    ### Climate (full) APRX layout and lists of mapframe, picture, text elements, and seasonal maps
    CLIMATE_FULL_APRX_LAYOUT_0 = CLIMATE_FULL_APRX.listLayouts()[0]
    CLIMATE_FULL_APRX_LAYOUT_0_PICTURE_ELEMENTS = CLIMATE_FULL_APRX_LAYOUT_0.listElements("picture_element")
    CLIMATE_FULL_APRX_LAYOUT_0_MAPFRAME_ELEMENTS = CLIMATE_FULL_APRX_LAYOUT_0.listElements("mapframe_element")
    CLIMATE_FULL_APRX_LAYOUT_0_TEXT_ELEMENTS = CLIMATE_FULL_APRX_LAYOUT_0.listElements("text_element")
    CLIMATE_FULL_APRX_OVERALL_MAP_0 = CLIMATE_FULL_APRX.listMaps('Overall')[0]
    CLIMATE_FULL_APRX_SPRING_MAP_0 = CLIMATE_FULL_APRX.listMaps('Spring')[0]
    CLIMATE_FULL_APRX_FALL_MAP_0 = CLIMATE_FULL_APRX.listMaps('Fall')[0]
    CLIMATE_FULL_APRX_SUMMER_MAP_0 = CLIMATE_FULL_APRX.listMaps('Summer')[0]
    CLIMATE_FULL_APRX_WINTER_MAP_0 = CLIMATE_FULL_APRX.listMaps('Winter')[0]

    ### Climate (no-overall) APRX layout and lists of mapframe, picture, text elements, and seasonal maps
    CLIMATE_NO_OVERALL_APRX_LAYOUT_0 = CLIMATE_NO_OVERALL_APRX.listLayouts()[0]
    CLIMATE_NO_OVERALL_APRX_LAYOUT_0_PICTURE_ELEMENTS = CLIMATE_NO_OVERALL_APRX_LAYOUT_0.listElements("picture_element")
    CLIMATE_NO_OVERALL_APRX_LAYOUT_0_MAPFRAME_ELEMENTS = CLIMATE_NO_OVERALL_APRX_LAYOUT_0.listElements("mapframe_element")
    CLIMATE_NO_OVERALL_APRX_LAYOUT_0_TEXT_ELEMENTS = CLIMATE_NO_OVERALL_APRX_LAYOUT_0.listElements("text_element")
    CLIMATE_NO_OVERALL_APRX_OVERALL_MAP_0 = CLIMATE_NO_OVERALL_APRX.listMaps('Overall')[0]
    CLIMATE_NO_OVERALL_APRX_SPRING_MAP_0 = CLIMATE_NO_OVERALL_APRX.listMaps('Spring')[0]
    CLIMATE_NO_OVERALL_APRX_FALL_MAP_0 = CLIMATE_NO_OVERALL_APRX.listMaps('Fall')[0]
    CLIMATE_NO_OVERALL_APRX_SUMMER_MAP_0 = CLIMATE_NO_OVERALL_APRX.listMaps('Summer')[0]
    CLIMATE_NO_OVERALL_APRX_WINTER_MAP_0 = CLIMATE_NO_OVERALL_APRX.listMaps('Winter')[0]

    # name of layer subset
    TMP_LAYER_SUBSET = INPUT_GDB_PATH+"\\test_lyr" 

    Z_VALUE = 5000

    ###### FUNCTION DEFINITIONS -- functions for title pages ######

    def updateTitleDoc(forest, forestsDict, pathToTitles, docxPath, overwrite = False):
        ##update forest name, region, and date in word doc and convert to pdf to create first 2 pages of TCA report

        outputName = os.path.join(pathToTitles,'TCA_TitlePages_{}.docx'.format(forest.replace(' ','_')))

        if os.path.exists(outputName) == overwrite or os.path.exists(outputName) == False: 

            document = Document(docxPath)

            region = 'Region '+str(forestsDict[forest])
        
            internal = document.paragraphs[2]

            txt = internal.text[:45]
            internal.text = ''
        
            obj_styles = document.styles
            obj_charstyle = obj_styles.add_style('CommentsStyle4', WD_STYLE_TYPE.CHARACTER)
            obj_font = obj_charstyle.font
            obj_font.size = Pt(12)
            obj_font.name = 'Arial'
            obj_font.bold=True
            obj_font.italic=True
            internal.add_run(txt, style = 'CommentsStyle4')
        
            obj_styles = document.styles
            obj_charstyle = obj_styles.add_style('CommentsStyle3', WD_STYLE_TYPE.CHARACTER)
            obj_font = obj_charstyle.font
            obj_font.size = Pt(12)
            obj_font.name = 'Arial'
            date = calendar.month_name[datetime.datetime.now().month] + ' '+str(datetime.datetime.now().year)
            internal.add_run(date, style = 'CommentsStyle3')
        
            section = document.paragraphs[5]       
            fontsize = 36 
            section.text = ''
        
            obj_styles = document.styles
            obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
            obj_font = obj_charstyle.font
            obj_font.size = Pt(fontsize)
            obj_font.name = 'Arial'
            section.add_run(forest, style = 'CommentsStyle')
        
            section2 = document.paragraphs[6]
            fontsize = 12
            section2.text = ''
            obj_charstyle2 = obj_styles.add_style('CommentsStyle2', WD_STYLE_TYPE.CHARACTER)
            obj_font2 = obj_charstyle2.font
            obj_font2.italic = True
            obj_font2.size = Pt(fontsize)
            obj_font2.name = 'Arial'

            section2.add_run(region, style = 'CommentsStyle2')

            if os.path.exists(outputName):
                os.remove(outputName)
            document.save(outputName)
            pdfOutputLoc = outputName[0:-4] + "pdf"

            print(f"Output location for pdf: {pdfOutputLoc}")
            print(f"Output location for word doc: {outputName}\n")
            print ('*** Do not open any docx files until all have been converted (this will overwrite current pdfs) ***\n')
        
            convert(outputName, pdfOutputLoc)

            ### MEMORY MANAGEMENT INTERNAL VARIABLES
            del forest, forestsDict, pathToTitles, docxPath, overwrite, outputName, document, region, section, txt, date, internal, obj_styles, obj_charstyle, obj_font, section2,fontsize, obj_charstyle2, obj_font2, pdfOutputLoc
            gc.collect()

    def seasonalPlot(nf_name, indicator, spring, summer ,winter, fall, meanList, base, fs = 18):#fores,reg,natl is order of meanlist   
        """ create stacked bar plot by seasons (for climate data)"""
    
        colorAr_mpl = COLOR_ARR/255.
        outer_colors = colorAr_mpl    
        width = 0.5       # the width of the bars: can also be len(x) sequence    
        fig, ax = plt.subplots()    
        scale_ar = np.array([spring, summer , fall, winter])    #order: very good, good, moderate, poor, very poor, nd
        labels = SEASONAL_SUB_COLUMNS
        ind = np.arange(len(labels)) 
        plt.rc('axes', labelsize=fs)
        plt.rc('figure', titlesize=fs)

        ax.bar(ind,height = scale_ar[:,4],width= width,  label='Very Poor',color= outer_colors[4,:], 
               linewidth = 0, align = 'center',tick_label = labels)#p,r,c
        ax.bar(ind,height = scale_ar[:,3], width=width, bottom=scale_ar[:,4],  label='Poor',color= outer_colors[3,:], 
               linewidth = 0, align = 'center')#, tick_label = labels)
        ax.bar(ind,height = scale_ar[:,2], width=width, bottom=scale_ar[:,3]+scale_ar[:,4],  label='Moderate',color= outer_colors[2,:], 
               linewidth = 0, align = 'center')#, tick_label = labels)
        ax.bar(ind, height = scale_ar[:,1], width=width, bottom=scale_ar[:,2]+scale_ar[:,3]+scale_ar[:,4], label='Good',color= outer_colors[1,:], 
               linewidth = 0, align = 'center')#, tick_label = labels)
        ax.bar(ind,height = scale_ar[:,0], width=width, bottom=scale_ar[:,1]+scale_ar[:,2]+scale_ar[:,3]+scale_ar[:,4], 
               label='Very Good',color= outer_colors[0,:],linewidth = 0, align = 'center')#, tick_label = labels)
        ax.bar(ind,height = scale_ar[:,5], width=width, bottom=scale_ar[:,0]+scale_ar[:,1]+scale_ar[:,2]+scale_ar[:,3]+scale_ar[:,4], 
               label='Null',color= outer_colors[5,:], linewidth = 0, align = 'center')
    
        for tick in ax.xaxis.get_major_ticks():
            tick.label.set_fontsize(fs) 
        ax.set_ylabel('% Landscapes in\nCondition Class', size = fs)
        plt.ylim(top=100)
        plt.ylim(bottom = 0)
    
        color = 'slategray'
        color2='w'
        ax2 = ax.twinx()  
        ax2.set_ylabel('Mean Condition Score', color = color)
    
        ax2.set_ylim((-1.,1.))
        ax2.set_yticks(np.arange(-1, 1.5, .5))
        ax2.set_yticklabels(np.arange(-1, 1.5, .5),color=color, size = fs)
    
        ax.set_ylim((0,100))
        ax.set_yticks(np.arange(0,120,20))
        ax.set_yticklabels(np.arange(0,120,20), size = fs)

        ax2.axhline(y=meanList[0],linewidth=2,color=color2,xmin=0.045,xmax=0.20, linestyle = '-')#.225
        ax2.axhline(y=meanList[0],linewidth=2,color=color,xmin=0.0456,xmax=0.1755, linestyle = '--') #.2292
    
        ax2.axhline(y=meanList[1],linewidth=2,color=color2,xmin=0.305,xmax=0.47, linestyle = '-')
        ax2.axhline(y=meanList[1],linewidth=2,color=color,xmin=0.305,xmax=0.435, linestyle = '--')
    
        ax2.axhline(y=meanList[2],linewidth=2,color=color2,xmin=0.55,xmax=0.69, linestyle = '-')
        ax2.axhline(y=meanList[2],linewidth=2,color=color,xmin=0.5645,xmax=0.695, linestyle = '--')
    
        ax2.axhline(y=meanList[3],linewidth=2,color=color2,xmin=0.81,xmax=0.97, linestyle = '-')
        ax2.axhline(y=meanList[3],linewidth=2,color=color,xmin=0.825,xmax=0.9545, linestyle = '--')

        plt.tick_params(
            axis='x',          # changes apply to the x-axis
            which='both',      # both major and minor ticks are affected
            bottom=False,      # ticks along the bottom edge are offfontdict={'fontsize': 14}
            top=False
            ,labelbottom=True
            ,labeltop = True
            ) # labels along the bottom edge are off
    
        outPth = os.path.join(base, nf_name)
        if os.path.exists(outPth) == False:
            os.mkdir(outPth)
        outFig = os.path.join(outPth, nf_name.replace(' ','_')+'_'+indicator.replace(' ','_')+'_figure.png')

        plt.savefig(outFig,dpi = 300,bbox_inches=BBOX_INCHES)

        ### MEMORY MANAGEMENT AFTER SAVE
        fig.clear()
        plt.close(fig) 
        ax.cla()
        ax2.cla()
        plt.cla()
        plt.clf()
        plt.close("all")
        del colorAr_mpl, outer_colors, width, scale_ar, labels, ind, ax2
        gc.collect()

    ###### FUNCTION DEFINITIONS -- functions for creating stacked bar plot with Nulls and creating final summary table ######
    def getPlot(nf_name, indicator, vals_nf,vals_reg, vals_natl, meanList, base, fs = 18):#fores,reg,natl is order of meanlist   
        """ create stacked bar plot """
    
        colorAr_mpl = COLOR_ARR/255.
        outer_colors = colorAr_mpl    
        width = 0.5       # the width of the bars: can also be len(x) sequence    
        fig, ax = plt.subplots()    
        scale_ar = np.array([vals_nf, vals_reg,vals_natl])    #order: very good, good, moderate, poor, very poor, nd
        labels = SPATIAL_SUB_COLUMNS
        ind = np.arange(len(labels)) 
        plt.rc('axes', labelsize=fs)
        plt.rc('figure', titlesize=fs)

        ax.bar(ind,height = scale_ar[:,4],width= width,  label='Very Poor',color= outer_colors[4,:], 
               linewidth = 0, align = 'center',tick_label = labels)#p,r,c
        ax.bar(ind,height = scale_ar[:,3], width=width, bottom=scale_ar[:,4],  label='Poor',color= outer_colors[3,:], 
               linewidth = 0, align = 'center')#, tick_label = labels)
        ax.bar(ind,height = scale_ar[:,2], width=width, bottom=scale_ar[:,3]+scale_ar[:,4],  label='Moderate',color= outer_colors[2,:], 
               linewidth = 0, align = 'center')#, tick_label = labels)
        ax.bar(ind, height = scale_ar[:,1], width=width, bottom=scale_ar[:,2]+scale_ar[:,3]+scale_ar[:,4], label='Good',color= outer_colors[1,:], 
               linewidth = 0, align = 'center')#, tick_label = labels)
        ax.bar(ind,height = scale_ar[:,0], width=width, bottom=scale_ar[:,1]+scale_ar[:,2]+scale_ar[:,3]+scale_ar[:,4], 
               label='Very Good',color= outer_colors[0,:],linewidth = 0, align = 'center')#, tick_label = labels)
        ax.bar(ind,height = scale_ar[:,5], width=width, bottom=scale_ar[:,0]+scale_ar[:,1]+scale_ar[:,2]+scale_ar[:,3]+scale_ar[:,4], 
               label='Null',color= outer_colors[5,:], linewidth = 0, align = 'center')
    
        for tick in ax.xaxis.get_major_ticks():
            tick.label.set_fontsize(fs) 
        ax.set_ylabel('% Landscapes in\nCondition Class', size = fs)
        plt.ylim(top=100)
        plt.ylim(bottom = 0)
    
        color = 'slategray'#'tab:grey'
        color2='w'
        ax2 = ax.twinx()  
        ax2.set_ylabel('Mean Condition Score', color = color)
    
        ax2.set_ylim((-1.,1.))
        ax2.set_yticks(np.arange(-1, 1.5, .5))
        ax2.set_yticklabels(np.arange(-1, 1.5, .5),color=color, size = fs)
    
        ax.set_ylim((0,100))
        ax.set_yticks(np.arange(0,120,20))
        ax.set_yticklabels(np.arange(0,120,20), size = fs)

        ax2.axhline(y=meanList[0],linewidth=2,color=color2,xmin=0.045,xmax=0.225, linestyle = '-')
        ax2.axhline(y=meanList[0],linewidth=2,color=color,xmin=0.0456,xmax=0.2292, linestyle = '--')
    
        ax2.axhline(y=meanList[1],linewidth=2,color=color2,xmin=0.41,xmax=0.599, linestyle = '-')
        ax2.axhline(y=meanList[1],linewidth=2,color=color,xmin=0.409,xmax=0.59595, linestyle = '--')
    
        ax2.axhline(y=meanList[2],linewidth=2,color=color2,xmin=0.77,xmax=0.97, linestyle = '-')
        ax2.axhline(y=meanList[2],linewidth=2,color=color,xmin=0.773,xmax=0.959, linestyle = '--')
        #ax2.scatter([0,1,2],meanList, color = 'none', marker = matplotlib.markers.HLINE, edgecolors = 'white')
        #ax2.tick_params(axis='y',labelcolor=color)
        plt.tick_params(
            axis='x',          # changes apply to the x-axis
            which='both',      # both major and minor ticks are affected
            bottom=False,      # ticks along the bottom edge are offfontdict={'fontsize': 14}
            top=False
            ,labelbottom=True
            ,labeltop = True
            ) # labels along the bottom edge are off
    
        outPth = os.path.join(base, nf_name)
        if os.path.exists(outPth) == False:
            os.mkdir(outPth)
        outFig = os.path.join(outPth, nf_name.replace(' ','_')+'_'+indicator.replace(' ','_')+'_figure.png')

        plt.savefig(outFig,dpi = 300,bbox_inches=BBOX_INCHES)

        #MEMORY MANAGEMENT OF FIGURES AND AXES AFTER SAVE AND INTERNAL VARIABLES
        fig.clear() 
        plt.close(fig) 
        fig.clf()
        ax.cla()
        ax2.cla() 
        plt.close("all") 
        plt.cla()
        plt.clf()
        del colorAr_mpl, outer_colors, width, scale_ar, labels, ind, ax2, outPth, outFig
        gc.collect()
    
    def buildIndicatorArray(indicator, unit, fields, gdb, shape):
        """
        build arrays for an indicator for a given unit (either region or forest)
    
        "unit" is either the name of a national forest or the number of a region
        """
        if type(unit) == str: #'National Forest' in 
            rowIndex = 0
        else: #region scale
            rowIndex = 1
        idx = fields.index(indicator)
        fields_cp = fields.copy()
        #fields_cp.append(indicator[:-6]) #add indicator name from attribute table minus '_class'
        fields_cp.append(indicator.replace('MC','M')) 

        fields_cp.append('TCA_Acres')
        indicatorCategoricalValues = []
        storeMeans = []
        areaValues = {'Very Good':0, 'Good':0, 'Moderate':0,'Poor':0,'Very Poor':0}

        with arcpy.da.SearchCursor(gdb, fields_cp) as cursor:
            for row in cursor:

                match = False

                #match either the Forest, or the region
                if type(unit) == str: 
                    if row[rowIndex] == unit:
                        match = True
                elif type(unit) == int: 
                    if int(row[rowIndex]) == unit:
                        match = True

                #if there is a match on either the forest or the region, then add to the 
                if match == True:

                    area = row[shape+1]
                    #area = row[shape]
                    if area is not None and row[idx] is not None and row[idx]!=None:
                        areaValues[row[idx]]+=float(area)

                    indicatorCategoricalValues.append(row[idx])

                    if row[idx] is not None: #

                        indicator_value = float(row[len(fields)])

                        ############# VALIDATION : Verifies that the indicator value that is being used to find the forest or regional mean is within the range of -1 to 1.
                        assert (indicator_value <= 1.0 and indicator_value >= -1.0), "GENERATE_REPORTS.py ln. 518. Indicator value for {0} in {1} found with value not within the expected range of (-1.0 to 1.0). Please check that the proper column is being referred to in the geodatabase, or that the values in the geodatabase are correct.".format(indicator, gdb)
                        #############

                        #this should be grabbing the actual calculated EMDS value to find the means per scale unit
                        storeMeans.append(indicator_value)

        counterArray = np.array([len([i for i in indicatorCategoricalValues if i == 'Very Good']), 
                                 len([i for i in indicatorCategoricalValues if i == 'Good']),
                                 len([i for i in indicatorCategoricalValues if i == 'Moderate']),
                                 len([i for i in indicatorCategoricalValues if i == 'Poor']),
                                 len([i for i in indicatorCategoricalValues if i == 'Very Poor']),
                                 len([i for i in indicatorCategoricalValues if i is None])])
        #default is nan, not 0.
        mean = np.nan

        if len(storeMeans) > 0:

            mean = np.mean(storeMeans)

            ############# VALIDATION : Verifies that the indicator value that is being used to find the forest or regional mean is within the range of -1 to 1.
            assert (mean <= 1.0 and mean >= -1.0), "GENERATE_REPORTS.py ln. 538. Calculated mean value for {0} in {1} not within the expected range of (-1.0 to 1.0). Please check that the proper indicator column is being referred to in the geodatase, or that the values in the geodatabase are correct.".format(indicator, gdb)
            #############

        ### MEMORY MANAGEMENT OF INTERNAL VARIABLES
        del idx, fields_cp, row, cursor, indicatorCategoricalValues, storeMeans
        gc.collect()

        return counterArray, mean, areaValues

    def buildNationalArray(indicator, fields, gdb):    
        idx = fields.index(indicator)
        fields_cp = fields.copy()

        ### include the calculated indicator value column
        fields_cp.append(indicator.replace('MC','M')) 

        fields_cp.append('TCA_Acres')
        shape = len(fields_cp)-1
        areaArray = {'Very Good':0, 'Good':0, 'Moderate':0,'Poor':0,'Very Poor':0}
        indicatorCategoricalValues = []
        storeMeans = []

        ### Loop throught the ...KPI_Assessment.gdb, and grab the values per indicator per LTA, within the national scale, to calculate the mean value for that indicator.
        with arcpy.da.SearchCursor(gdb, fields_cp) as cursor:

            for row in cursor:

                area = row[shape]
                if area is not None and area != None and type(area) == float and row[idx] is not None and row[idx] != None:
                    areaArray[row[idx]] += float(area)
                indicatorCategoricalValues.append(row[idx])
                if row[idx] is not None:

                    indicator_value = float(row[len(fields)])

                    ############# VALIDATION : Verifies that the indicator value that is being used to find the forest or regional mean is within the range of -1 to 1.
                    assert (indicator_value <= 1.0 and indicator_value >= -1.0), "GENERATE_REPORTS.py ln. 574. Indicator value for {0} in {1} found with value not within the expected range of (-1.0 to 1.0). Please check that the proper column is being referred to in the geodatabase, or that the values in the geodatabase are correct.".format(indicator, gdb)
                    #############

                    storeMeans.append(indicator_value)

        counterArray = np.array([len([i for i in indicatorCategoricalValues if i == 'Very Good']),
                                 len([i for i in indicatorCategoricalValues if i == 'Good']),
                                 len([i for i in indicatorCategoricalValues if i == 'Moderate']), 
                                 len([i for i in indicatorCategoricalValues if i == 'Poor']),
                                 len([i for i in indicatorCategoricalValues if i == 'Very Poor']),
                                 len([i for i in indicatorCategoricalValues if i is None])
                                 ])
        mean = np.nan

        if len(storeMeans) > 0:

            mean = np.mean(storeMeans)

            ############# VALIDATION : Verifies that the indicator value that is being used to find the forest or regional mean is within the range of -1 to 1.
            assert (mean <= 1.0 and mean >= -1.0), "GENERATE_REPORTS.py ln. 593. Calculated mean value for {0} in {1} not within the expected range of (-1.0 to 1.0). Please check that the proper indicator column is being referred to in the geodatase, or that the values in the geodatabase are correct.".format(indicator, gdb)
            #############

        ### MEMORY MANAGEMENT OF INTERNAL VARIABLES
        del idx, fields_cp, shape, indicatorCategoricalValues, storeMeans, cursor, row
        gc.collect()

        return counterArray, mean, areaArray

    def getArray(keyList, fields, listIndicators,gdb):     
        forestsArrayDict = {}
        forestsMeanDict = {}
        forestsAreaDict = {}
        for forest in keyList:
            tempArrayDict = {}
            tempMeans = {}
            tempAreaDict ={}
            for indicator in fields[2:]:
                ar, mean, area = buildIndicatorArray(indicator, forest, fields, gdb, len(listIndicators)+2)
                commonName = [k for k in listIndicators.keys() if listIndicators[k]==indicator][0]
                tempArrayDict[commonName]=ar 
                tempMeans[commonName] = mean
                tempAreaDict[commonName] = area
            forestsArrayDict[forest] = tempArrayDict
            forestsMeanDict[forest] = tempMeans
            forestsAreaDict[forest] = tempAreaDict

        return forestsArrayDict,forestsMeanDict, forestsAreaDict

    def toPercent(ar):

        ar_sum = sum(ar) #total number of LTAs in given unit'
        #ar = ini_array[::-1] #reverse
        ar_pct  = np.array([i/ar_sum*100 for i in ar])
        return ar_pct

    def render_mpl_table_final(outPath, data):
        
        size = (np.array(data.shape[::-1]) + np.array([MPL_FINAL_COL_WIDTH, 2])) * np.array([MPL_FINAL_COL_WIDTH, MPL_FINAL_ROW_HEIGHT])
        fig, ax = plt.subplots(figsize=size)
        ax.axis('off')

        mpl_table = ax.table(cellText=data.values, bbox=BBOX, colLabels=data.columns,rowLabels= data.index,loc='top')

        mpl_table.auto_set_font_size(False)
        mpl_table.set_fontsize(FONT_SIZE)
    
        for k, cell in mpl_table._cells.items():
            cell.set_edgecolor(EDGE_COLOR)
            if k[0] == 11 or k[0] == 6 or k[0] == 0:# or k[0] == 9 or k[0] == 5:
                    h = cell.get_height()
                    cell.set_height(h*1.5)
            if k[0] == 0 or k[1] < HEADER_COLUMNS:
                cell.set_text_props(weight='bold', color='w')
                cell.set_facecolor(HEADER_COLOR)                      
            else:
                cell.set_facecolor(ROW_COLORS[k[0]%len(ROW_COLORS)])   

        plt.savefig(outPath, dpi = 300, bbox_inches=BBOX_INCHES)

        ### MEMORY MANAGEMENT OF FIGURES AND AXES AFTER SAVE AND INTERNAL VARIABLES
        fig.clear()
        plt.close(fig) 
        ax.cla()
        fig.clf()
        plt.close("all")
        plt.cla()
        plt.clf()
        del outPath, data, size, mpl_table, k, cell, fig, ax
        gc.collect()

    def createFinalTable(forestDict, regionDict, nationDict, indicList, nf_name, base):
        indicatorList = list(indicList)    #WHY? this is already a list of indicators.wouldn't this be a list of lists?
    
        indicatorList.sort()    
        indicatorList.remove('Overall Terrestrial Condition')
        indicatorList.insert(0,'Overall Terrestrial Condition')

        vstack = np.zeros((len(indicatorList),5))#forestDict[indicatorList[0]])
    
        for row in range(len(indicatorList)):
            col=-1
            for this in [forestDict[indicatorList[row]],regionDict[indicatorList[row]],nationDict[indicatorList[row]],
                             forestDict[indicatorList[row]] - regionDict[indicatorList[row]], 
                             forestDict[indicatorList[row]] - nationDict[indicatorList[row]]]:
                col+=1
                vstack[row,col] = round(this,4)
    
        indicatorList = ['Forest Insect and Pathogen\nHazard' if 'Insect' in x else x for x in indicatorList]
        indicatorList = ['Wildfire Hazard Potential' if 'Wildfire Potential' in x else x for x in indicatorList]
        indicatorList = ['Uncharacteristic Disturbance\nEvents' if 'Uncharacteristic Disturbance' in x else x for x in indicatorList]
    
        df = pd.DataFrame(vstack, index = indicatorList, columns = ['Forest Mean','Region Mean', 'Nation Mean', 'Departure\nFrom Region','Departure\nFrom Nation'])
        df.fillna("", inplace=True)

        outPath = os.path.join(base, nf_name.replace(' ','_')+'_finalTable.png')

        render_mpl_table_final(outPath, df)
    
        del forestDict, regionDict, nationDict, indicList, nf_name, base, indicatorList, vstack, df, outPath
        gc.collect()

    ###### FUNCTION DEFINITIONS -- functions for creating tables ######
    def render_mpl_table(outPath, data, heading):

        #if ax is None:
        size = (np.array(data.shape[::-1]) + np.array([MPL_COL_WIDTH, 1])) * np.array([MPL_COL_WIDTH, MPL_ROW_HEIGHT])
        fig, ax = plt.subplots(figsize=size)
        ax.axis('off')
    
        mpl_table_header = ax.table(cellText=[['']],colLabels = [heading],loc='top')
    
        mpl_table = ax.table(cellText=data.values, bbox=BBOX, colLabels=data.columns,rowLabels= data.index,loc='top')
        #ax.axhline(y=1, color=header_color, lw=1,alpha=1)  #alpha=1, 
        ax.axvline(x=0,color='w', alpha=1, lw=1)
        ax.axvline(x=0,color='w', alpha=1, lw=1)

        mpl_table.auto_set_font_size(False)
        mpl_table.set_fontsize(FONT_SIZE)
        mpl_table_header.auto_set_font_size(False)
        mpl_table_header.set_fontsize(FONT_SIZE)
    
        for k, cell in mpl_table._cells.items():
            cell.set_edgecolor(EDGE_COLOR)
            if k[0] == 0 or k[1] < HEADER_COLUMNS:
                cell.set_text_props(weight='bold', color='w')
                cell.set_facecolor(HEADER_COLOR)
            
                if k[0] == 0:
                    cell.set_text_props(weight='bold', color='w')#, va ='top')
                
            else:
                cell.set_facecolor(ROW_COLORS[k[0]%len(ROW_COLORS) ])
            
        for k, cell in mpl_table_header._cells.items():
            cell.set_edgecolor(EDGE_COLOR)
            if k[0] == 0 or k[1] < HEADER_COLUMNS:           
                cell.set_facecolor(HEADER_COLOR)
                #h = cell.get_height()
                #cell.set_height(h/1.5)
                cell.set_text_props(weight='bold', color='w')#, verticalalignment ='bottom')
        
        #adjust height of first 'row' of initial table
        cellDict = mpl_table_header.get_celld()
    
        cellDict[(1,0)].set_height(0)
        if len(heading) > 25:
            cellDict[(0,0)].set_height(1/3.8)#6.) 
        else:
            cellDict[(0,0)].set_height(1/6.) 

        # print("attempting to save with outpath : {}".format(outPath))
        plt.savefig(outPath, dpi=300,bbox_inches=BBOX_INCHES)

        #MEMORY MANGEMENT
        fig.clear() 
        plt.close(fig) 
        ax.cla()
        fig.clf()
        plt.close("all")
        plt.cla()
        plt.clf()
        del outPath, data, heading, fig, ax
        gc.collect()
    

    def conv2Class(gdb, classField):
        arcpy.management.AddField(gdb,classField,'TEXT')
        valueField = classField.replace("_class",'')
        arcpy.management.CalculateField(gdb, "{}".format(classField), "fun(!{}!)".format(valueField), "PYTHON3", """def fun(a):
            if a <= -.6:
                return 'Very Poor'
            if a > -.6 and a <= -.2:
                return 'Poor'
            if a > -.2 and a<= 0.2:
                return 'Moderate'
            if a >.2 and a <=.6:
                return 'Good'
            else:
                return 'Very Good'""", "TEXT")

    def exportAndAppendToFinalPDF(srcLayout, tempPath, finalPath):
        srcLayout.exportToPDF(tempPath)#,resolution=500, output_as_image=True) #output as image to convert vector parts of pdf to raster (helps with width issue of LTA outline)
        finalPath.appendPages(tempPath)


    ######## PRIMARY SCRIPT #########

    if __name__ == 'MODULES.GENERATE_CONUS_REPORTS':

        ### Build output directory
        if os.path.exists(FINAL_OUTPUT_PATH):
            if os.path.isdir(FINAL_OUTPUT_PATH):
                if os.path.exists(OUTPUT_TITLE_PAGES_PATH):
                    pass
                else:
                    os.mkdir(OUTPUT_TITLE_PAGES_PATH)
        else:
            #build from scratch
            os.mkdir(FINAL_OUTPUT_PATH)
            os.mkdir(OUTPUT_TITLE_PAGES_PATH)

        ### CHECK/ADD attribute fields not in gdb
        for i in list(CONUS_EMDS_INDICATORS_FOR_REPORTS.values()):
            if i not in FIELD_NAMES_IN_GDB:
                print('error {} not in fields in above dict'.format(i))
                conv2Class(INPUT_GDB_FEATURE_CLASS_PATH, i)
                #ensure correct attribute names NOTE: must also have correct attr names for forestname and region below in fields list

        UPDATED_FIELDS = list(EXTRA_CONUS_FIELDS.values()) + list(CONUS_EMDS_INDICATORS_FOR_REPORTS.values())

        ### build national array using dictionary of array
        nationalArrayDict = {}
        nationalMeansDict = {}
        nationalAreaDict = {}

        for indicator in UPDATED_FIELDS[2:]:
            ar, natl_mean, area = buildNationalArray(indicator, UPDATED_FIELDS, INPUT_GDB_FEATURE_CLASS_PATH)
            commonName = [k for k in CONUS_EMDS_INDICATORS_FOR_REPORTS.keys() if CONUS_EMDS_INDICATORS_FOR_REPORTS[k]==indicator][0]
            nationalArrayDict[commonName] = ar
            nationalMeansDict[commonName] = natl_mean
            nationalAreaDict[commonName] = area

        forestsArrayDict, forestMeansDict, forestsAreaDict = getArray(CONUS_FORESTS_AND_REGIONS.keys(), UPDATED_FIELDS, CONUS_EMDS_INDICATORS_FOR_REPORTS, INPUT_GDB_FEATURE_CLASS_PATH)
        regionalArrayDict, regionalMeansDict, regionalAreaDict = getArray(FS_REGIONS, UPDATED_FIELDS, CONUS_EMDS_INDICATORS_FOR_REPORTS, INPUT_GDB_FEATURE_CLASS_PATH)

        print("Building plots and tables for all CONUS indicators, for all CONUS forests. Please continue to be patient...\n")

        ### iterate through forests, and build figures + tables
        for forest in FOREST_KEYS_LIST:

            output_forest_path = os.path.join(FINAL_OUTPUT_PATH,forest) 

            if os.path.exists(output_forest_path) ==False:
                os.mkdir(output_forest_path)

            updateTitleDoc(forest, CONUS_FORESTS_AND_REGIONS, OUTPUT_TITLE_PAGES_PATH, TITLE_PAGES_TEMPLATE, True) #last param is whether to overwrite or not
    
            region = CONUS_FORESTS_AND_REGIONS[forest]

            createFinalTable(forestMeansDict[forest], regionalMeansDict[region], nationalMeansDict, INDICATOR_KEYS_LIST, forest, output_forest_path)
    
            ### iterate through indicators to create tables/graphs
            for indicator in INDICATOR_KEYS_LIST:
                meanList = [forestMeansDict[forest][indicator], regionalMeansDict[region][indicator], nationalMeansDict[indicator]]
                #make sure they're listed from very poor (first in array) to very good (last)
                vals_nf = toPercent(forestsArrayDict[forest][indicator])
                vals_reg = toPercent(regionalArrayDict[region][indicator])
                vals_natl = toPercent(nationalArrayDict[indicator])
            
                getPlot(forest, indicator,vals_nf,vals_reg, vals_natl, meanList, FINAL_OUTPUT_PATH)
           
                ### CREATE TABLES
                forest_ar = [int(i/1000.) for i in forestsAreaDict[forest][indicator].values()]
                region_ar = [int(i/1000.) for i in regionalAreaDict[region][indicator].values()]
                nation_ar = [int(i/1000.) for i in nationalAreaDict[indicator].values()]
        
                ### Arrays continaing number of LTAs by indicator per forest or per region or nationally
                forest_ar_counter = forestsArrayDict[forest][indicator][:-1]
                region_ar_counter = regionalArrayDict[region][indicator][:-1]
                nation_ar_counter = nationalArrayDict[indicator][:-1]
        
                ### build table for acres
                full_ar = np.array([forest_ar,region_ar, nation_ar]).T #5 rows, 3 cols
            
                ### need to create unit col names to apply lambda function to add commas and convert # to int
                df = pd.DataFrame(full_ar, index=TCA_CLASSIFICATIONS, columns=SPATIAL_SUB_COLUMNS)
            
                for column in df:
                    df[column] = df.apply(lambda x: "{:,}".format(x[column]), axis=1)

                output_acre_table_path = os.path.join(output_forest_path,(forest+'_'+indicator+'_AcreTable.png').replace(' ','_'))
            
                render_mpl_table(output_acre_table_path, df, LANDSCAPE_THOUSAND_ACRES_TABLE_TITLE)
            
                ### build table for Number
                full_ar = np.array([forest_ar_counter,region_ar_counter, nation_ar_counter]).T #5 rows, 3 cols
            
                ### need to create unit col names to apply lambda function to add commas and convert # to int
                df = pd.DataFrame(full_ar, index=TCA_CLASSIFICATIONS, columns=SPATIAL_SUB_COLUMNS)
            
                for column in df:            
                    df[column] = df.apply(lambda x: "{:,}".format(x[column]), axis=1)

                output_count_table_path = os.path.join(output_forest_path,(forest+'_'+indicator+'_CountTable.png').replace(' ','_'))
            
                render_mpl_table(output_count_table_path, df, NUMBER_OF_LANDSCAPES_TABLE_TITLE)
            
                del indicator, meanList, vals_nf, vals_reg, forest_ar, region_ar, nation_ar, forest_ar_counter, region_ar_counter, nation_ar_counter, full_ar, df, output_acre_table_path, output_count_table_path
                gc.collect()

            ### run climate figures and recreate final tables
            for climIndicator in CONUS_CLIMATE_INDICATOR_CATEGORIES:
                listSeasonsForInd = [season + ' ' + climIndicator for season in SEASONAL_SUB_COLUMNS]
                meanList = [forestMeansDict[forest][x] for x in listSeasonsForInd]
                spring = toPercent(forestsArrayDict[forest][listSeasonsForInd[0]])
                summer = toPercent(forestsArrayDict[forest][listSeasonsForInd[1]])
                fall = toPercent(forestsArrayDict[forest][listSeasonsForInd[2]])
                winter = toPercent(forestsArrayDict[forest][listSeasonsForInd[3]])

                seasonalPlot(forest, climIndicator, spring, summer ,winter, fall, meanList, FINAL_OUTPUT_PATH, fs = 18)
            
                spring_ar = [int(i/1000.) for i in forestsAreaDict[forest][listSeasonsForInd[0]].values()]
                summer_ar = [int(i/1000.) for i in forestsAreaDict[forest][listSeasonsForInd[1]].values()]
                fall_ar = [int(i/1000.) for i in forestsAreaDict[forest][listSeasonsForInd[2]].values()]
                winter_ar = [int(i/1000.) for i in forestsAreaDict[forest][listSeasonsForInd[3]].values()]
        
                ### Arrays continaing number of LTAs by indicator per forest by season
                spring_ar_counter = forestsArrayDict[forest][listSeasonsForInd[0]][:-1]
                summer_ar_counter = forestsArrayDict[forest][listSeasonsForInd[1]][:-1]
                fall_ar_counter = forestsArrayDict[forest][listSeasonsForInd[2]][:-1]
                winter_ar_counter = forestsArrayDict[forest][listSeasonsForInd[3]][:-1]
        
                ### build table for acres
                full_ar = np.array([spring_ar,summer_ar, fall_ar, winter_ar]).T #5 rows, 3 cols
            
                ### need to create unit col names to apply lambda function to add commas and convert # to int
                df = pd.DataFrame(full_ar,index=TCA_CLASSIFICATIONS, columns=SEASONAL_SUB_COLUMNS)
                for column in df:
                    df[column] = df.apply(lambda x: "{:,}".format(x[column]), axis=1)

                output_acre_table_path = os.path.join(output_forest_path,(forest+'_'+climIndicator+'_AcreTable.png').replace(' ','_'))
            
                render_mpl_table(output_acre_table_path, df, LANDSCAPE_THOUSAND_ACRES_TABLE_TITLE)  
            
                ### build table for Number
                full_ar = np.array([spring_ar_counter,summer_ar_counter, fall_ar_counter, winter_ar_counter]).T #5 rows, 3 cols
            
                ### need to create unit col names to apply lambda function to add commas and convert # to int
                df = pd.DataFrame(full_ar,index=TCA_CLASSIFICATIONS, columns=SEASONAL_SUB_COLUMNS)
                for column in df:            
                    df[column] = df.apply(lambda x: "{:,}".format(x[column]), axis=1)

                output_count_table_path = os.path.join(output_forest_path,(forest+'_'+climIndicator+'_CountTable.png').replace(' ','_'))
            
                render_mpl_table(output_count_table_path, df, NUMBER_OF_LANDSCAPES_TABLE_TITLE)
            
                del listSeasonsForInd, meanList, spring, summer, fall, winter, spring_ar, summer_ar, fall_ar, winter_ar, spring_ar_counter, summer_ar_counter, fall_ar_counter, winter_ar_counter, full_ar, df, output_acre_table_path, output_count_table_path
                gc.collect()
            
            ### MANAGE MEMORY OUTER LOOP
            del output_forest_path, region
            gc.collect()
    
            pageNumber=2
            regionNum = CONUS_FORESTS_AND_REGIONS[forest]
            PDF_path = os.path.join(FINAL_OUTPUT_PATH,"TCA_{a}_Summary_R{c}_{b}.pdf".format(a=TCA_VERSION.replace('TCA.',''),c=regionNum, b=forest.replace(' ','_')))  

            #Create the final PDF file and append title pages
            if os.path.exists(PDF_path):
                os.remove(PDF_path)

            ### PDF that will have all forest-level data and pages appended
            finalPDF = arcpy.mp.PDFDocumentCreate(PDF_path)

            titlePage = os.path.join(OUTPUT_TITLE_PAGES_PATH, 'TCA_TitlePages_'+forest.replace(' ','_')+".pdf")#.format(forest.replace(' ','_')))
            finalPDF.appendPages(titlePage)
    
            query_forest = EXTRA_CONUS_FIELDS['forest_name_column'] + " = " +"'"+forest+"'"      #query = "LTA_FS_Forest" +'=' +"'"+forest+"'"
    
            forests_subset = DISSOLVED_FOREST_LAYER_PATH + "_subset"

            queriedLayer_forest = arcpy.management.MakeFeatureLayer(DISSOLVED_FOREST_LAYER_PATH, forests_subset, query_forest).getOutput(0)
    
            arcpy.management.SelectLayerByAttribute(forests_subset, "CLEAR_SELECTION")  
    
            queriedLayer_forest.definitionQuery = query_forest#"LTA_FS_Forest = '"+forest+"'" #doesn't have symbology applied
    
            arcpy.management.SelectLayerByAttribute(queriedLayer_forest, "CLEAR_SELECTION")  

            # apply symbology
            arcpy.management.ApplySymbologyFromLayer(queriedLayer_forest, SYMBOLOGY_LAYER_FORESTS_PATH, "VALUE_FIELD OBJECTID OBJECTID") #lookback

            print("Constructing CONUS report pages, please continue to be patient...\n")

            #Non-climate related indicators. Each gets own page now.
            non_climate_indicators = []
            for indicator in CONUS_EMDS_INDICATORS_FOR_REPORTS.keys():
                if indicator not in CONUS_EMDS_CLIMATE_INDICATORS:
                    non_climate_indicators.append(indicator)
    
            for indicator in non_climate_indicators: 

                pageNumber +=1
        
                # create queried layer 
                queriedLayer = arcpy.management.MakeFeatureLayer(INPUT_GDB_FEATURE_CLASS_PATH,TMP_LAYER_SUBSET, query_forest).getOutput(0)
                #clear selection on layer
                arcpy.management.SelectLayerByAttribute(TMP_LAYER_SUBSET, "CLEAR_SELECTION")  
        
                ## apply query
                queriedLayer.definitionQuery = query_forest
                #clear selection
                arcpy.management.SelectLayerByAttribute(queriedLayer, "CLEAR_SELECTION")  
                # get indicators
                indicator_field = CONUS_EMDS_INDICATORS_FOR_REPORTS[indicator]
        
                #add symbology properties 
                symbology_properties = "VALUE_FIELD {0} {1}".format(CONUS_EMDS_INDICATOR_TREE_ROOT, indicator_field.replace('MC','M'))
                
                arcpy.management.ApplySymbologyFromLayer(queriedLayer, SYMBOLOGY_LAYER_FINAL_PATH, symbology_properties)
        
                ########################################
        
                #add queried layer - LTA TCA indicator queried to NF - to both top and bottom maps 
                TOP_MAP.addLayer(queriedLayer, "TOP")
        
                # apply symbology agian (this may be redundant with above apply symbology from layer)
                arcpy.management.ApplySymbologyFromLayer(TOP_MAP.listLayers()[0], SYMBOLOGY_LAYER_FINAL_PATH, symbology_properties)#.getOutput(0) #lookback
        
                # make layers visible or not (only visible layers = basemap topography, top layer (queried TCA lTAs), outline of forest queried to this NF) 
                TOP_MAP.addLayer(queriedLayer_forest, "TOP")
        
                for lyr in TOP_MAP.listLayers()[2:]:
                    if lyr.name == 'World Terrain Base' or lyr.name == 'World Ocean Base' or lyr.name == 'World Imagery':
                        lyr.visible = True
                    else:
                        lyr.visible = False
        
                subset_extent = str(APRX_LAYOUT_0_MAPFRAME_0.getLayerExtent(TOP_MAP.listLayers()[0])).split()
        
                listNewExt = [float(subset_extent[0])-Z_VALUE,float(subset_extent[1])-Z_VALUE, float(subset_extent[2])+Z_VALUE, float(subset_extent[3])+Z_VALUE]
       
                APRX_LAYOUT_0_MAPFRAME_0.camera.setExtent(arcpy.Extent(listNewExt[0], listNewExt[1], listNewExt[2], listNewExt[3]))#, False, True)) #selection only, symbolized extent
       
                watermarkText = [i for i in APRX_LAYOUT_0_TEXT_ELEMENTS if 'WatermarkText' in i.name]
            
                for elm in APRX_LAYOUT_0_TEXT_ELEMENTS:
                    if elm.name =='PageNumber':
                        elm.text = str(pageNumber)
                    elif elm.name =='IndicatorSubtextTop':
                        elm.text = CONUS_EMDS_INDICTORS_SUBTEXT[indicator]
                    elif elm.name == 'TopTableText1':
                        elm.text = 'Comparison of Number of Landscapes by Condition Class'
                    elif elm.name == 'TopTableText2':
                        elm.text = 'Comparison of Landscape Acres (in thousands) by Condition Class'
                    elif elm.name == 'TopFigureTxt':
                        elm.text = 'Comparison of Mean Condition and Proportion of Landscapes in Condition Classes (*Gray refers to Landscapes with No Rating)'
                    elif elm.name == 'TopMapTitle':
                        elm.text = indicator
                        if indicator == 'Uncharacteristic Disturbance':
                            elm.text = 'Uncharacteristic Disturbance Events'
                        elm.elementPositionX =  0.5 
                    elif elm.name =='ForestedSystemsBanner':
                        elm.text = CONUS_FORESTED_INDICATOR_HEADERS[indicator]
                        if indicator == 'Grassland Encroachment':
                            if forest in CONUS_FORESTS_WITH_NO_NON_FORESTED_LANDTYPES:
                                watermarkText[0].text = 'No Non-Forested Landscapes'
                                watermarkText[0].visible = True
                            else:
                                watermarkText[0].visible = False
                        elif indicator == 'Grassland Productivity':
                            if forest in CONUS_FORESTS_WITH_NO_NON_FORESTED_LANDTYPES:
                                watermarkText[0].text = 'No Non-Forested Landscapes'
                                watermarkText[0].visible = True
                            else:
                                watermarkText[0].visible = False
                        elif indicator =='Tree Mortality':
                            if forest in CONUS_FORESTS_WITH_NO_FOREST_LANDTYPE:
                                watermarkText[0].text = 'No Forested Landscapess'
                                watermarkText[0].visible = True
                        else:
                            watermarkText[0].visible = False 
                     
                pngName = (forest+'_'+indicator+'_figure.png').replace(' ','_')
                tableName1 = (forest+'_'+indicator+'_CountTable.png').replace(' ','_')
                tableName2 = (forest+'_'+indicator+'_AcreTable.png').replace(' ','_')

                picturePath = os.path.join(FINAL_OUTPUT_PATH, forest, pngName)
                tablePath1_1 = os.path.join(FINAL_OUTPUT_PATH, forest, tableName1)
                tablePath1_2 = os.path.join(FINAL_OUTPUT_PATH, forest, tableName2)

                for pic in APRX_LAYOUT_0_PICTURE_ELEMENTS:
                    if pic.name == 'TopFigure':
                        pic.sourceImage = picturePath
                    elif pic.name == 'TableTop1':
                        pic.sourceImage = tablePath1_1
                    elif pic.name == 'TableTop2':
                        pic.sourceImage = tablePath1_2
                    elif pic.name =='USDAlogo':
                        pic.sourceImage = USDA_LOGO_PATH

                APRX.save()
        
                if os.path.exists(TEMP_PDF_PATH):
                    os.remove(TEMP_PDF_PATH)
            
                exportAndAppendToFinalPDF(APRX_LAYOUT_0, TEMP_PDF_PATH, finalPDF)
        
                #Commit changes and delete variable reference
                for ll in TOP_MAP.listLayers()[0:-2]:
                    if ll.name == TEST_LAYER_PATH or ll.name == SYMBOLOGY_LAYER_PATH or ll.name == 'tbl3' or ll.name == 'TCA_v2_LTA_Baseline_top':
                        TOP_MAP.removeLayer(ll)

                ########### MEMORY MANAGEMENT ############
                del pngName,  tableName1 
                del tableName2, picturePath, 
                del tablePath1_1, tablePath1_2
                del watermarkText
        
                gc.collect()
                ##########################################

            forestPicName = (forest+'_finalTable.png').replace(' ','_')

            forestTableSource = os.path.join(FINAL_OUTPUT_PATH, forest, forestPicName)
    
            for pic in APRX_FINAL_PAGE_LAYOUT_0_PICTURE_ELEMENTS:
                if pic.name == 'Picture':            
                    pic.sourceImage = forestTableSource        
    
            exportAndAppendToFinalPDF(APRX_FINAL_PAGE_LAYOUT_0, TEMP_PDF_PATH, finalPDF)

            pageNumber += 2 #add extra page number for mean condition table

            drought_map = DROUGHT_APRX.listMaps('TopPane')[0]

            indicator_field = CONUS_EMDS_INDICATORS_FOR_REPORTS['Drought']

            #add symbology properties 
            symbology_properties = "VALUE_FIELD {0} {1}".format(CONUS_EMDS_INDICATOR_TREE_ROOT, indicator_field.replace('MC','M')) #OLD VERS USED SYMBOLOGY FROM CLASS FIELD AKA RAW FIELD CLASSIFIED INTO POOR, GOOD, ETC.

            # apply symbology 
            arcpy.management.ApplySymbologyFromLayer(queriedLayer, SYMBOLOGY_LAYER_FINAL_PATH, symbology_properties)#.getOutput(0) #lookback
    
            #add queried layer - LTA TCA indicator queried to NF - to both top and bottom maps 
            drought_map.addLayer(queriedLayer, "TOP")
            arcpy.management.ApplySymbologyFromLayer(drought_map.listLayers()[0], SYMBOLOGY_LAYER_FINAL_PATH, symbology_properties)#.getOutput(0) #lookback
            drought_map.addLayer(queriedLayer_forest, "TOP")

            for lyr in drought_map.listLayers()[2:]:
                if lyr.name == 'World Terrain Base' or lyr.name == 'World Ocean Base' or lyr.name == 'World Imagery':
                    lyr.visible = True
                else:
                    lyr.visible = False
    
            subset_extent = str(DROUGHT_APRX_MAPFRAME_0.getLayerExtent(drought_map.listLayers()[0])).split()

            listNewExt = [float(subset_extent[0])-Z_VALUE,float(subset_extent[1])-Z_VALUE, float(subset_extent[2])+Z_VALUE, float(subset_extent[3])+Z_VALUE]

            DROUGHT_APRX_MAPFRAME_0.camera.setExtent(arcpy.Extent(listNewExt[0], listNewExt[1], listNewExt[2], listNewExt[3]))#, False, True)) #selection only, symbolized extent
    
            for elm in DROUGHT_APRX_LAYOUT_0_TEXT_ELEMENTS:
                if elm.name =='PageNumber':
                    elm.text = str(pageNumber)
                    break

            pngName = (forest+'_Drought_figure.png').replace(' ','_')
            tableName1 = (forest+'_Drought_CountTable.png').replace(' ','_')
            tableName2 = (forest+'_Drought_AcreTable.png').replace(' ','_')
    
            picturePath = os.path.join(FINAL_OUTPUT_PATH, forest, pngName)
            tablePath1_1 = os.path.join(FINAL_OUTPUT_PATH, forest, tableName1) #count table - num of las
            tablePath1_2 = os.path.join(FINAL_OUTPUT_PATH, forest, tableName2)

            pictureElements = DROUGHT_APRX_LAYOUT_0_PICTURE_ELEMENTS

            for pic in pictureElements:
                if pic.name =='BottomFigure':
                    pic.sourceImage = picturePath
                elif pic.name == 'TableBottom1': #number of landscapes
                    pic.sourceImage = tablePath1_1
                elif pic.name == 'TableBottom2':
                    pic.sourceImage = tablePath1_2

            if os.path.exists(TEMP_PDF_PATH):
                os.remove(TEMP_PDF_PATH)
        
            exportAndAppendToFinalPDF(DROUGHT_APRX_LAYOUT_0, TEMP_PDF_PATH, finalPDF)

            ### remove temporary layers from the drought map.
            for myMap in DROUGHT_APRX.listMaps():
                for ll in myMap.listLayers()[0:-4]:
                    if ll.name == TEST_LAYER_PATH or ll.name == SYMBOLOGY_LAYER_PATH:
                        myMap.removeLayer(ll)
                
            ######## MEMORY MANAGEMENT ########
            del forestPicName, forestTableSource, pngName
            del drought_map, tableName1, tableName2, picturePath, indicator_field, symbology_properties
            del tablePath1_1, tablePath1_2, pictureElements, subset_extent, listNewExt
            gc.collect()
            ###################################
    
            for climVar in CONUS_CLIMATE_APPENDIX_PAGES:

                if 'Percent' in climVar:
                    _climate_aprx = CLIMATE_NO_OVERALL_APRX
                    _climate_aprx_layout_0 = CLIMATE_NO_OVERALL_APRX_LAYOUT_0
                    overall = CLIMATE_NO_OVERALL_APRX_OVERALL_MAP_0
                    spring = CLIMATE_NO_OVERALL_APRX_SPRING_MAP_0
                    fall = CLIMATE_NO_OVERALL_APRX_FALL_MAP_0
                    summer = CLIMATE_NO_OVERALL_APRX_SUMMER_MAP_0
                    winter = CLIMATE_NO_OVERALL_APRX_WINTER_MAP_0
                    _picture_elements = CLIMATE_NO_OVERALL_APRX_LAYOUT_0_PICTURE_ELEMENTS
                    _mapframe_elements = CLIMATE_NO_OVERALL_APRX_LAYOUT_0_MAPFRAME_ELEMENTS
                    _textElements = CLIMATE_NO_OVERALL_APRX_LAYOUT_0_TEXT_ELEMENTS
                else:
                    _climate_aprx = CLIMATE_FULL_APRX
                    _climate_aprx_layout_0 = CLIMATE_FULL_APRX_LAYOUT_0
                    overall = CLIMATE_FULL_APRX_OVERALL_MAP_0
                    spring = CLIMATE_FULL_APRX_SPRING_MAP_0
                    fall = CLIMATE_FULL_APRX_FALL_MAP_0
                    summer = CLIMATE_FULL_APRX_SUMMER_MAP_0
                    winter = CLIMATE_FULL_APRX_WINTER_MAP_0
                    _picture_elements = CLIMATE_FULL_APRX_LAYOUT_0_PICTURE_ELEMENTS
                    _mapframe_elements = CLIMATE_FULL_APRX_LAYOUT_0_MAPFRAME_ELEMENTS
                    _textElements = CLIMATE_FULL_APRX_LAYOUT_0_TEXT_ELEMENTS
            
                pageNumber +=1 
                tag = climVar.replace('Overall ', '') 
                listRelevant = [c for c in CONUS_EMDS_CLIMATE_INDICATORS if c.endswith(tag)]#.append('Overall '+tag)    
        
                for idxName in listRelevant:
                    season = idxName.replace(' ' + tag, '')
                    indicator_field = CONUS_EMDS_INDICATORS_FOR_REPORTS[idxName]
                    #add symbology properties 
                    symbology_properties = "VALUE_FIELD {0} {1}".format(CONUS_EMDS_INDICATOR_TREE_ROOT, indicator_field.replace('MC','M')) #OLD VERS USED SYMBOLOGY FROM CLASS FIELD AKA RAW FIELD CLASSIFIED INTO POOR, GOOD, ETC. #lookback
                    # apply symbology 
                    arcpy.management.ApplySymbologyFromLayer(queriedLayer, SYMBOLOGY_LAYER_FINAL_PATH, symbology_properties)#.getOutput(0) #lookback
            
                    thisMap = _climate_aprx.listMaps(season)[0]
            
                    #add queried layer - LTA TCA indicator queried to NF - to both top and bottom maps 
                    thisMap.addLayer(queriedLayer, "TOP")
                    arcpy.management.ApplySymbologyFromLayer(thisMap.listLayers()[0], SYMBOLOGY_LAYER_FINAL_PATH, symbology_properties)#.getOutput(0) #lookback
                    thisMap.addLayer(queriedLayer_forest, "TOP")

                    for lyr in thisMap.listLayers()[2:]:
                        if lyr.name == 'World Terrain Base' or lyr.name == 'World Ocean Base' or lyr.name == 'World Imagery':
                            lyr.visible = True
                        else:
                            lyr.visible = False

                for mapframeElement in _mapframe_elements:
                    thisMap = _climate_aprx.listMaps(mapframeElement.name.replace('_MapFrame',''))[0]    
                    subset_extent = str(mapframeElement.getLayerExtent(thisMap.listLayers()[0])).split()
                    listNewExt = [float(subset_extent[0])-Z_VALUE,float(subset_extent[1])-Z_VALUE, float(subset_extent[2])+Z_VALUE, float(subset_extent[3])+Z_VALUE]
                    mapframeElement.camera.setExtent(arcpy.Extent(listNewExt[0], listNewExt[1], listNewExt[2], listNewExt[3]))#, False, True)) #selection only, symbolized extent

                ### updates each of the text elements in the layout
                for elm in _textElements:  
                    if elm.name =='PageNumber':
                        elm.text = str(pageNumber)
                    elif elm.name =='IndicatorSubtextTop':
                        elm.text = CONUS_EMDS_INDICTORS_SUBTEXT[climVar]                
                    elif '_Header' in elm.name and 'Percent' not in climVar:
                        season = elm.name.replace('_Header','')
                        elm.text = season + climVar.replace('Overall','')
                    elif elm.name == 'TopMapTitle':
                        elm.text = climVar
                        elm.elementPositionX = 0.5 #+ elm.elementWidth/2       
                
                    elif elm.name == 'TopTableText1':
                        elm.text = 'Comparison of Number of Landscapes by Condition Class by Season'

                    elif elm.name == 'TopTableText2':
                        elm.text = 'Comparison of Landscape Acres (in thousands) by Condition Class by Season'
                         
                    elif elm.name == 'TopFigureTxt':
                        elm.text = 'Comparison of Mean Condition and Proportion of Landscapes in Condition Class by Season'
                    
                    elif elm.name == 'TopMapTitle':
                        elm.text = climVar.replace('Overall Precipitation Exposure', 'Overall Precipitation Exposure')  # removed "-Amount" suffix from precip var      
                        elm.elementPositionX =  0.5 #+ elm.elementWidth/2    
                    elif elm.name == 'WatermarkText':
                        if climVar == 'Overall Precipitation Exposure Percent': # removed "-Percent Change" suffix from precip var on left of dict and added "Percent"
                            elm.visible = True
                        else:
                            elm.visible = False
                    
                pngName = (forest + '_' + CONUS_FIGURE_PATH_DICT[climVar] + '_figure.png').replace(' ','_')
                tableName1 = (forest + '_'+ CONUS_FIGURE_PATH_DICT[climVar] + '_CountTable.png').replace(' ','_')
                tableName2 = (forest + '_' + CONUS_FIGURE_PATH_DICT[climVar] + '_AcreTable.png').replace(' ','_')

                picturePath = os.path.join(FINAL_OUTPUT_PATH, forest, pngName)
                tablePath1_1 = os.path.join(FINAL_OUTPUT_PATH, forest, tableName1)
                tablePath1_2 = os.path.join(FINAL_OUTPUT_PATH, forest, tableName2)
        
                ### assigns the chart or table paths to the sources per element in the list of picture elements
                for pic in _picture_elements:
                    if pic.name == 'TopFigure':
                        pic.sourceImage = picturePath
                    elif pic.name == 'TableTop1':
                        pic.sourceImage = tablePath1_1
                    elif pic.name == 'TableTop2':
                        pic.sourceImage = tablePath1_2
        
                if os.path.exists(TEMP_PDF_PATH):
                    os.remove(TEMP_PDF_PATH)

                exportAndAppendToFinalPDF(_climate_aprx_layout_0, TEMP_PDF_PATH, finalPDF)

                ###moved this into this loop
                for myMap in _climate_aprx.listMaps():
                    for ll in myMap.listLayers()[0:-2]:
                        if ll.name == TEST_LAYER_PATH or ll.name == SYMBOLOGY_LAYER_PATH:
                            myMap.removeLayer(ll) 

                ###### MEMORY MANAGEMENT #######
                del climVar, _climate_aprx, _climate_aprx_layout_0, overall, spring, fall, summer, winter, _mapframe_elements, _picture_elements, _textElements 
                del idxName, pngName, pic, tableName1, tableName2, picturePath, tablePath1_1, tablePath1_2, 
                del mapframeElement, thisMap, subset_extent, listNewExt, symbology_properties, season, listRelevant, tag

                gc.collect()
                ################################

            ### append the appendix and glossary pages and save.
            finalPDF.appendPages(APPENDIX_PDF_PATH)
            finalPDF.appendPages(GLOSSARY_PDF_PATH)    
            finalPDF.saveAndClose()

            print(f"\nFinal Report for {forest} complete : {PDF_path}\n")

            ###### MEMORY MANAGEMENT #######
            arcpy.management.ClearWorkspaceCache()
            del queriedLayer, lyr, finalPDF, regionNum, PDF_path, titlePage, query_forest, queriedLayer_forest
            gc.collect()
            ################################

        #remove temp file if exists
        if os.path.exists(TEMP_PDF_PATH):
            os.remove(TEMP_PDF_PATH)

        print("CONUS reports complete. Please continue to be patient...")
